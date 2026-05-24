"""
Week 2 보고서 v3 — 피드백 반영 최종본
수정 사항:
  1) 4.2절: 센서 배치 출처를 "데이터셋 문서 기반" → "상관분석 결과 바탕 추정"으로 변경
  2) PS3: 1차 회로 단정 → PS1과 음의 상관, PS5와 양의 상관 → 추가 확인 필요 명시
  3) Week 4: MFPCA/Transformer/Hierarchical → 향후 연구 방향으로만 언급, 현실적 축소
  4) 2.3절: Valve 결과표 추가
  5) AI 문체 완화
서식:
  - 맑은 고딕 전체, 본문 11pt, 소제목 13pt 굵게, 대제목 16pt 굵게
  - 그림 설명: 맑은 고딕 11pt, 기울임 없음
  - 그래프: 페이지당 최대 2개
  - ~하였다 체
"""
import os, shutil
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.join(os.path.dirname(__file__), '..')
FIG_DIR = os.path.join(BASE, 'reports', 'figures')
OUT_PATH = os.path.join(BASE, 'reports', '2주차_인턴십_보고서.docx')
OUT_COPY = r'C:\Users\pjeon\OneDrive\바탕 화면\과제 및 레포트\2주차_인턴십_보고서_v5.docx'

doc = Document()

FONT = '맑은 고딕'
SZ_BODY = Pt(11)
SZ_H1 = Pt(16)
SZ_H2 = Pt(13)
SZ_H3 = Pt(11)
SZ_CAP = Pt(11)   # 그림 캡션도 11pt

# ── 글꼴 설정 헬퍼 ──
def sf(run, size=SZ_BODY, bold=False, color=None):
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    run.font.italic = False
    if color:
        run.font.color.rgb = color
    rpr = run.element.find(qn('w:rPr'))
    if rpr is None:
        rpr = OxmlElement('w:rPr')
        run.element.insert(0, rpr)
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.insert(0, rf)
    rf.set(qn('w:eastAsia'), FONT)

# ── 스타일 ──
style = doc.styles['Normal']
style.font.name = FONT; style.font.size = SZ_BODY
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
style.paragraph_format.line_spacing = 1.15

for lv, sz in [(1, SZ_H1), (2, SZ_H2), (3, SZ_H3)]:
    h = doc.styles[f'Heading {lv}']
    h.font.name = FONT; h.font.size = sz; h.font.bold = True
    h.font.italic = False; h.font.color.rgb = RGBColor(0,0,0)
    h.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

# ── 본문/불릿/표/그림 헬퍼 ──
def body(text):
    p = doc.add_paragraph(); r = p.add_run(text); sf(r); return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    for rn in p.runs: rn.text = ''
    r = p.add_run(text); sf(r); return p

def table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); sf(r, Pt(10), bold=True)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i+1].cells[j]; c.text = ''
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(v)); sf(r, Pt(10))
    doc.add_paragraph()
    return t

def fig(fname, caption, width=Inches(5.5)):
    path = os.path.join(FIG_DIR, fname)
    if not os.path.exists(path):
        body(f'[그림 누락: {fname}]'); return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = p.add_run(); rn.add_picture(path, width=width)
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(caption); sf(r, SZ_CAP)  # 11pt, 기울임 없음
    doc.add_paragraph()

def pb(): doc.add_page_break()


# ═══════════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════════
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run('2주차 인턴십 진행 보고서'); sf(r, Pt(22), bold=True)

p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(12)
r2 = p2.add_run('유압 시스템 다중결함 동시진단 (UCI Hydraulic Systems)')
sf(r2, Pt(13), color=RGBColor(80,80,80))

p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p3.paragraph_format.space_before = Pt(20)
r3 = p3.add_run('2026-05-23  |  박정태 (202221215)  |  부산대 기계공학부  |  노유정 교수님 랩')
sf(r3, Pt(9), color=RGBColor(100,100,100))

pl = doc.add_paragraph(); pl.paragraph_format.space_before = Pt(6)
rl = pl.add_run('━' * 55); sf(rl, Pt(8), color=RGBColor(150,150,150))


# ═══════════════════════════════════════════════
# 개요
# ═══════════════════════════════════════════════
doc.add_heading('2주차 목표 및 개요', level=1)

body(
    '2주차의 핵심 목표는 1주차에서 구축한 EDA 기반 위에 '
    '(1) Helwig et al.(2015) 논문의 LDA baseline을 정밀 재현하고, '
    '(2) 다양한 ML 모델과 불균형 처리 기법을 비교하며, '
    '(3) 1주차에서 발견한 PS4 센서의 독자적 거동에 대해 물리적 원인을 규명하는 것이었다. '
    '이하 모든 실험은 StandardScaler 정규화, 5-fold Stratified CV 기준으로 평가하였다.'
)
body(
    '또한 1주차 보고서에 대한 교수님 피드백을 반영하여, '
    'PS4의 물리적 원인을 정량적 분석으로 규명하고, '
    '멀티 샘플링 레이트 입력 전략과 향후 계획을 본 보고서에 포함하였다. '
    '2주차에는 LDA baseline 재현, 불균형 처리 실험, PS4 물리적 원인 규명을 완료하였고, '
    'ML 5종 비교와 멀티 샘플링 레이트 입력 전략을 추가로 수행하였다.'
)


# ═══════════════════════════════════════════════
# 1. LDA Baseline
# ═══════════════════════════════════════════════
pb()
doc.add_heading('1. LDA Baseline 정밀 재현', level=1)

doc.add_heading('1.1 실험 설정', level=2)
body(
    'Helwig et al.(2015)의 접근법에 기반하여 11종의 시간도메인 통계량을 구성하였다. '
    '17개 센서에서 사이클별로 mean, std, rms, peak, crest factor, '
    'skewness, kurtosis, peak-to-peak, shape factor, impulse factor, clearance factor를 '
    '추출하여 187차원 특성 벡터를 만들었다. '
    '정규화 후 LDA를 적용하였다.'
)

doc.add_heading('1.2 결과', level=2)
table(
    ['Component', 'Classes', 'Accuracy (%)', 'Macro-F1 (%)', 'MCC', '재현 기준 (%)', '달성'],
    [
        ['Cooler',      '3', '99.95 (±0.09)', '99.93 (±0.12)', '0.9992', '99.0', 'O'],
        ['Valve',       '4', '99.86 (±0.18)', '99.83 (±0.22)', '0.9980', '98.0', 'O'],
        ['Pump',        '3', '99.77 (±0.14)', '99.70 (±0.20)', '0.9963', '95.0', 'O'],
        ['Accumulator', '4', '98.19 (±0.20)', '97.93 (±0.25)', '0.9741', '98.0', 'O'],
    ]
)
body(
    '4개 컴포넌트 모두 Helwig et al.(2015)의 보고치를 재현 기준으로 설정하였으며, 전부 달성하였다. '
    'Cooler(99.95%)와 Valve(99.86%)는 거의 완벽한 분류를 보였다. '
    'Accumulator(98.19%)가 상대적으로 가장 낮았는데, '
    '어큐뮬레이터의 가스 챔버가 압력 맥동을 완충하여 열화 정도가 1차 회로 센서에 간접적으로만 나타나기 때문으로 보았다. '
    '실제로 Confusion Matrix(그림 1-4)를 보면, '
    '130 클래스를 115(severely reduced) 클래스로 오분류한 경우가 15건으로 가장 빈번하게 발생하였으며, '
    '이는 두 상태의 신호 특성이 유사함을 반영한다.'
)

doc.add_heading('1.3 Confusion Matrix', level=2)
fig('lda_cm_cooler.png',
    '[그림 1-1] LDA Confusion Matrix — Cooler (Accuracy: 99.95%)')
fig('lda_cm_valve.png',
    '[그림 1-2] LDA Confusion Matrix — Valve (Accuracy: 99.86%)')

pb()
fig('lda_cm_pump.png',
    '[그림 1-3] LDA Confusion Matrix — Pump (Accuracy: 99.77%)')
fig('lda_cm_accumulator.png',
    '[그림 1-4] LDA Confusion Matrix — Accumulator (Accuracy: 98.19%)')


# ═══════════════════════════════════════════════
# 2. 불균형 데이터 처리 실험
# ═══════════════════════════════════════════════
pb()
doc.add_heading('2. 불균형 데이터 처리 실험', level=1)

doc.add_heading('2.1 실험 배경', level=2)
body(
    'UCI Hydraulic 데이터셋에서는 Valve(optimal 클래스가 약 3.1배), '
    'Pump(no_leakage가 약 2.5배) 등의 클래스 불균형이 존재하였다. '
    '이러한 불균형이 소수 클래스의 recall에 미치는 영향을 확인하기 위해 '
    'Random Forest 분류기를 기준으로 4종 실험을 설계하였다.'
)

doc.add_heading('2.2 실험 설계', level=2)
table(
    ['실험', '방법', '설명'],
    [
        ['E1', 'Baseline', 'Random Forest (n=100), 아무 처리 없음'],
        ['E2', 'class_weight', 'RF + class_weight="balanced"'],
        ['E3', 'SMOTE', 'RF + SMOTE 오버샘플링'],
        ['E4', 'SMOTE+weight', 'RF + SMOTE + class_weight="balanced"'],
    ]
)
body('평가 지표: Accuracy, Macro-F1, MCC, 클래스별 Recall')

doc.add_heading('2.3 주요 결과', level=2)

# ── Valve 결과표 추가 (피드백 반영) ──
doc.add_heading('Valve (4클래스: optimal / small lag / severe lag / close to failure)', level=3)
table(
    ['실험', 'Accuracy (%)', 'Macro-F1 (%)', 'Min Recall (%)'],
    [
        ['E1: Baseline',      '99.05', '98.73', '97.50'],
        ['E2: class_weight',  '98.91', '98.55', '97.22'],
        ['E3: SMOTE',         '98.96', '98.59', '97.78'],
        ['E4: SMOTE+weight',  '98.96', '98.59', '97.78'],
    ]
)

doc.add_heading('Pump (3클래스: no leakage / weak / severe)', level=3)
table(
    ['실험', 'Accuracy (%)', 'Macro-F1 (%)', 'Min Recall (%)'],
    [
        ['E1: Baseline',      '99.50', '99.30', '98.98'],
        ['E2: class_weight',  '99.59', '99.43', '99.19'],
        ['E3: SMOTE',         '99.64', '99.50', '99.39'],
        ['E4: SMOTE+weight',  '99.64', '99.50', '99.39'],
    ]
)

doc.add_heading('Accumulator (4클래스: optimal / slightly reduced / severely reduced / close to failure)', level=3)
table(
    ['실험', 'Accuracy (%)', 'Macro-F1 (%)', 'Min Recall (%)'],
    [
        ['E1: Baseline',      '98.50', '98.27', '96.24'],
        ['E2: class_weight',  '98.78', '98.64', '96.99'],
        ['E3: SMOTE',         '98.78', '98.62', '96.74'],
        ['E4: SMOTE+weight',  '98.78', '98.62', '96.74'],
    ]
)

body(
    'Cooler는 3클래스가 비교적 균형적이어서 4종 실험 간 차이가 거의 없었으므로 '
    '별도 표를 생략하였다.'
)

fig('imbalance_macro_f1.png',
    '[그림 2-1] 불균형 처리 4종 실험 Macro-F1 비교')

doc.add_heading('2.4 분석', level=2)
body(
    '전반적으로 불균형 처리의 효과가 크지 않았다. '
    '187차원 시간도메인 특성의 분리력이 이미 충분하여, '
    '클래스 불균형이 있어도 분류기가 소수 클래스를 잘 구분하였기 때문이다.'
)
bullet('Valve: 3.1배 불균형이 있었으나, Baseline(E1)에서 이미 Macro-F1 98.73%로 높았다. '
       'class_weight(E2) 적용 시 98.55%로 소폭 변화하였으나, '
       '이는 CV의 측정 불확실성 범위 내로 통계적으로 유의미하지 않을 수 있다. '
       'RF에서 class_weight는 분기 시 불순도 가중치에 영향을 주는데, '
       '이미 분리가 명확한 경우 소수 클래스 과대 가중이 다수 클래스 정밀도를 소폭 희생시킬 수 있다.')
bullet('Pump: SMOTE(E3)가 Min Recall을 98.98%에서 99.39%로 개선하였다 (+0.41%p). '
       'Macro-F1도 99.30%에서 99.50%로 상승하였다.')
bullet('Accumulator: class_weight(E2)가 Min Recall을 96.24%에서 96.99%로, '
       'Macro-F1을 98.27%에서 98.64%로 개선하였다.')
bullet(
    '실험 결과, 187차원 통계 특성의 분리력이 이미 높아 Valve와 Pump에서는 불균형 처리 효과가 거의 없었다. '
    '다만 4클래스 분류인 Accumulator에서는 class_weight 적용 시 Min Recall이 96.24%에서 96.99%로 개선되어, '
    '클래스 경계가 복잡할수록 처리 효과가 나타남을 확인하였다.'
)


# ═══════════════════════════════════════════════
# 3. ML Baseline 5종 비교
# ═══════════════════════════════════════════════
pb()
doc.add_heading('3. ML Baseline 5종 비교', level=1)

doc.add_heading('3.1 실험 설정', level=2)
body(
    'LDA에 더해 4종의 앙상블 모델(Random Forest, XGBoost, LightGBM, Extra Trees)을 비교하였다. '
    '모든 모델에 n_estimators=200을 적용하였으며, 앞서 서술한 기본 실험 설정(StandardScaler 정규화 및 5-fold Stratified CV)과 동일하게 유지하였다.'
)

doc.add_heading('3.2 결과', level=2)
table(
    ['Component, Macro-F1 (%)', 'LDA', 'RF', 'XGBoost', 'LightGBM', 'ExtraTrees'],
    [
        ['Cooler',       '99.9', '99.9', '99.8', '99.9', '99.9'],
        ['Valve',        '99.8', '98.8', '98.9', '98.9', '98.7'],
        ['Pump',         '99.7', '99.4', '99.6', '99.8', '99.3'],
        ['Accumulator',  '97.9', '98.6', '98.4', '98.8', '98.9'],
    ]
)

fig('ml_baseline_comparison.png',
    '[그림 3-1] ML Baseline 5종 모델 Macro-F1 비교 (5-Fold CV)')

doc.add_heading('3.3 분석', level=2)
body('5개 모델 모두 Macro-F1 97.9% 이상의 높은 성능을 보였다.')
bullet(
    'Cooler, Valve: LDA가 가장 높았다(각각 99.9%, 99.8%). '
    '선형 분리가 잘 되는 경우에는 복잡한 모델이 추가 이점을 주지 못하였다.'
)
bullet('Pump: LightGBM(99.8%)이 LDA(99.7%)를 약간 상회하였다.')
bullet(
    'Accumulator: ExtraTrees(98.9%)가 LDA(97.9%)보다 1.0%p 높았다. '
    '4클래스 분류에서 비선형 경계를 다루는 데 앙상블 모델이 유리하였다.'
)
bullet(
    '전체적으로 모델 간 차이가 크지 않았다. '
    '이는 187차원 통계 특성의 분리력이 이미 높다는 의미이다. '
    '따라서 3주차 딥러닝의 목표는 추가 정확도 향상보다, '
    '수작업 특성 추출 과정을 없애고 원시 시계열에서 자동으로 패턴을 학습할 수 있는지 확인하는 데 두었다. '
    '이는 새로운 결함 유형이 발생했을 때 특성 재설계 없이 모델을 재학습할 수 있다는 실용적 이점과 연결된다.'
)


# ═══════════════════════════════════════════════
# 4. PS4 물리적 원인 검증 (수정: 출처 표현 + PS3 불일치)
# ═══════════════════════════════════════════════
pb()
doc.add_heading('4. PS4 센서 독자적 거동의 물리적 원인 검증', level=1)

body(
    '※ 1주차 보고서 교수님 피드백: '
    '"PS4의 독자적 거동에 대해 데이터셋 관련 정보를 찾아 물리적 원인을 명확히 할 것"'
)

doc.add_heading('4.1 배경', level=2)
body(
    '1주차 EDA에서 PS4(압력 센서 4번)가 다른 압력 센서들과 다른 거동을 보이는 것을 발견하였다. '
    '이에 대해 3가지 가설을 세우고 정량적으로 검증하였다.'
)
bullet('가설 1: PS4가 어큐뮬레이터 분기에 위치하여 accumulator 라벨과 높은 상관을 보인다.')
bullet('가설 2: PS4가 2차 냉각/여과 회로에 위치하여 PS5/PS6과 높은 상관을 보인다.')
bullet('가설 3: PS4가 밸브 하류에 위치하여 valve 라벨과 높은 상관을 보인다.')

# ── 수정: 센서 배치를 "상관분석 결과 바탕 추정"으로 변경 ──
doc.add_heading('4.2 시스템 구성 및 센서 배치 추정', level=2)
body(
    'UCI 데이터셋 설명에 따르면, 본 유압 시스템은 '
    '1차 작동 회로(primary working circuit)와 '
    '2차 냉각/여과 회로(secondary cooling-filtration circuit)로 구성되어 있다. '
    '그러나 데이터셋 문서에는 각 센서의 정확한 설치 위치가 명시되어 있지 않다.'
)
body(
    '따라서 아래의 센서 회로 분류는 4.4절의 피어슨 상관계수 분석 결과를 바탕으로 추정한 것이며, '
    '확정적 정보가 아님을 밝힌다.'
)
bullet('PS1, PS2: 1차 작동 회로 소속으로 추정')
bullet('PS4, PS5, PS6: 2차 냉각/여과 회로 소속으로 추정')
bullet('PS3: 1차 회로와 2차 회로 양쪽의 영향을 모두 받는 것으로 보이며, 정확한 위치는 추가 확인이 필요')

body(
    '각 추정의 근거가 되는 구체적인 상관계수 수치와 해석은 4.4절에서 상세히 다룬다.'
)

doc.add_heading('4.3 검증 1: Mutual Information 분석', level=2)
body(
    '각 압력 센서의 시간도메인 특성과 4개 결함 라벨 간의 '
    'Mutual Information(MI)을 계산하였다.'
)
fig('ps4_mutual_information.png',
    '[그림 4-1] 압력 센서별 결함 라벨 Mutual Information 히트맵')

body('주요 결과는 다음과 같다:')
bullet('PS4-accumulator MI = 0.1255로, PS1-accumulator MI(0.3435)의 약 37% 수준이었다.')
bullet('PS4는 accumulator보다 cooler 라벨과 더 높은 MI를 보였다.')
bullet('따라서 가설 1(어큐뮬레이터 분기)은 기각하였다.')

doc.add_heading('4.4 검증 2: 피어슨 상관계수 분석', level=2)
body(
    '6개 압력 센서의 사이클별 평균값으로 피어슨 상관계수를 계산하였다.'
)
fig('ps4_correlation_matrix.png',
    '[그림 4-2] 압력 센서 간 피어슨 상관계수 행렬')

body('PS4와 다른 센서 간 상관계수는 다음과 같다 (그림 4-2 참조):')
bullet('PS4-PS1: r = 0.043 (거의 독립적)')
bullet('PS4-PS2: r = -0.020 (거의 독립적)')
bullet('PS4-PS3: r = 0.484 (중간 수준의 양의 상관)')
bullet('PS4-PS5: r = 0.745 (강한 양의 상관)')
bullet('PS4-PS6: r = 0.745 (강한 양의 상관)')
body(
    'PS4는 PS1-PS2(1차 회로)와 거의 무관하였고, PS5-PS6과 강한 양의 상관을 보였다. '
    '이 결과는 가설 2(2차 냉각/여과 회로)를 지지하였다.'
)
body(
    '한편, PS1-PS2는 서로 r=0.995의 매우 강한 양의 상관을 보여 동일 회로(1차 회로)에 속하는 것으로 추정하였으며, '
    'PS5-PS6은 r=1.000으로 사실상 동일한 거동을 나타내 같은 회로로 판단하였다. '
    'PS3의 경우, PS1과 r=-0.720(강한 음의 상관), PS5와 r=0.755(강한 양의 상관)를 보였다. '
    'PS1-PS2(1차 회로)와는 반대 방향으로 움직이면서 PS4-PS5(2차 회로)와는 같은 방향이므로, '
    '단순히 1차 회로로 분류하기 어려웠다. '
    'PS3은 두 회로의 분기 지점에 설치되어 양쪽 영향을 모두 받는 것으로 추정하였으며, '
    '정확한 위치는 시스템 배관도(P&ID) 없이는 확정할 수 없다.'
)

pb()
doc.add_heading('4.5 검증 3: PSD 분석', level=2)
fig('ps4_psd_comparison.png',
    '[그림 4-3] 압력 센서 PSD 비교 (PS1-PS3 vs PS4)')

body(
    'PSD 분석 결과에서도 PS1-PS2는 0.4Hz 부근에 최대 에너지가 집중된 유사한 형태를 보인 반면, '
    'PS4는 11.7Hz 부근에 별도의 피크가 나타나는 등 스펙트럼 형태 자체가 달랐다. '
    '이는 PS4가 물리적으로 다른 위치에 있음을 지지하는 결과이다.'
)

fig('ps4_psd_by_accumulator.png',
    '[그림 4-4] PS4 PSD — Accumulator 상태별 비교')

body(
    'Accumulator 상태별로 PS4의 PSD를 비교하면, '
    '열화 정도에 따른 스펙트럼 변화가 관찰되기는 하였으나, '
    '이는 시스템 전체 압력 맥동의 간접적 영향으로 해석하였다.'
)

doc.add_heading('4.6 종합 결론', level=2)
body(
    'MI, 상관계수, PSD 세 가지 분석을 종합한 결과, '
    'PS4는 2차 냉각/여과 회로에 위치한 센서로 추정하였다. '
    'PS1-PS2와 거의 독립적이었으며, PS5와 r=0.745의 강한 상관을 보였다. '
    '1주차에 관찰한 PS4의 독자적 거동은 PS4가 다른 유압 회로에 연결되어 있기 때문이었다.'
)
body(
    '다만 센서의 정확한 배치를 확정하기 위해서는 데이터셋 원저자의 시스템 배관도(P&ID)가 필요하며, '
    '본 분석은 상관분석 결과에 기반한 추정임을 한계로 밝힌다.'
)
body(
    '상관분석 결과가 2차 회로 가설을 지지하므로, '
    '배관도 확인 전이라도 회로별 센서 그룹화를 딥러닝 branch 설계의 출발점으로 삼는 것은 타당하다고 판단하였다. '
    '다만 추후 P&ID 확인을 통해 branch 구성을 검증할 필요가 있다.'
)


# ═══════════════════════════════════════════════
# 5. 멀티 샘플링 레이트 입력 전략
# ═══════════════════════════════════════════════
pb()
doc.add_heading('5. 멀티 샘플링 레이트 입력 전략', level=1)

body(
    '※ 1주차 보고서 교수님 피드백: '
    '"멀티 샘플링 레이트를 학습 모델에 어떻게 입력할지 방안을 고민할 것"'
)

doc.add_heading('5.1 현황', level=2)
body('UCI Hydraulic 데이터셋의 17개 센서는 3가지 샘플링 레이트로 수집되었다:')
table(
    ['샘플링 레이트', '센서', '개수', '사이클당 샘플 수'],
    [
        ['100 Hz', 'PS1-PS6 (압력), EPS1 (모터 전력)', '7개', '6,000'],
        ['10 Hz',  'FS1, FS2 (유량)', '2개', '600'],
        ['1 Hz',   'TS1-TS4 (온도), VS1 (진동), CE (효율), CP (냉각 전력), SE (안정성)', '8개', '60'],
    ]
)
body(
    '2주차까지는 사이클별 통계량(187차원)을 추출하여 센서 간 샘플링 레이트의 불일치 문제를 해결하였으나, '
    '3주차에서 원시 시계열 데이터를 딥러닝 모델에 직접 입력하기 위해서는 별도의 입력 전략이 요구된다.'
)

doc.add_heading('5.2 검토 전략', level=2)

body(
    '검토한 세 전략은 다음과 같다.'
)
body(
    '전략 1(업샘플링 통일)은 모든 센서를 100Hz로 선형 보간하여 6,000 timesteps로 맞추는 방법이다. '
    '구현이 간단하지만, 1Hz 센서(VS1 등)를 100배 보간하면 실제로 없는 정보가 생기는 문제가 있다.'
)
body(
    '전략 2(Multi-Branch CNN)는 샘플링 레이트별로 별도 branch를 구성하는 방법이다. '
    '100Hz branch(7ch x 6000), 10Hz branch(2ch x 600), 1Hz branch(8ch x 60)로 나누어 '
    '각 branch에서 특성을 뽑은 뒤 합친다. 구현은 복잡해지지만 각 센서의 원래 해상도를 보존할 수 있다.'
)
body(
    '전략 3(통계 특성 + Raw 혼합)은 100Hz 센서만 raw 시계열로 CNN에 넣고, '
    '나머지 센서는 통계량을 뽑아 MLP에 넣은 뒤 두 경로를 합치는 방법이다.'
)
body(
    '각 센서 고유의 시간 해상도를 왜곡 없이 보존할 수 있는 전략 2를 우선적으로 시도하며, '
    '추후 4장에서 도출한 유압 회로별 도메인 지식을 해당 브랜치 내 채널 그룹화에 연계하여 확장할 계획이다. '
    '3주차에서는 PyTorch로 Multi-Branch CNN을 구현하여 전략 2를 검증하고, '
    '전략 1(100Hz 업샘플링)과의 성능 차이를 비교할 계획이다.'
)


# ═══════════════════════════════════════════════
# 6. 향후 계획 (수정: Week 4 현실적으로 축소)
# ═══════════════════════════════════════════════
pb()
doc.add_heading('6. 향후 계획 (1~4주차 전체 로드맵)', level=1)

body('※ 1주차 보고서 교수님 피드백: "향후 계획을 다음 자료에 같이 공유할 것"')
body('본 인턴십은 4주간 진행되며, 주차별 목표와 현황은 다음과 같다.')

doc.add_heading('6.1 전체 로드맵', level=2)
table(
    ['주차', '핵심 과제', '세부 내용', '상태'],
    [
        ['Week 1', 'EDA + LDA Baseline',
         '데이터셋 구조 파악, 17개 센서 시각화, 라벨 분포 분석, '
         '187차원 특성 추출, LDA 5-fold CV baseline',
         '완료'],
        ['Week 2', 'ML Baselines + 심화 분석',
         'LDA 정밀 재현, 불균형 4종 실험, ML 5종 비교, '
         'PS4 원인 검증, 멀티 샘플링 레이트 전략 수립',
         '완료'],
        ['Week 3', '딥러닝 모델 구현',
         'Multi-Scale CNN 구현, Multi-Task Learning '
         '(공유 backbone + 4개 head), '
         'Multi-Branch CNN (샘플링 레이트별), '
         'ML baseline과 성능 비교',
         '예정'],
        ['Week 4', '실험 완료 + 최종 정리',
         'Multi-Branch CNN 실험 마무리, '
         '전체 결과 비교 분석 (ML vs DL), '
         '최종 보고서 작성',
         '예정'],
    ]
)

doc.add_heading('6.2 연구 가설 현황', level=2)
table(
    ['가설', '내용', '검증 시기', '현황'],
    [
        ['H1', '압력센서 간 상관이 복합결함 커플링을 포착',
         'Week 2', 'PS4 회로 분리 확인 (부분 검증)'],
        ['H2', 'Multi-rate 센서 융합이 naive upsampling보다 유리',
         'Week 3', '전략 수립 완료, 구현 예정'],
        ['H4', 'Multi-task learning으로 4개 결함 동시 학습',
         'Week 3', '구현 예정'],
    ]
)
body(
    'H3(MFPCA), H5(Transformer attention)는 4주 일정 내에 구현하지 않기로 하고, '
    '이후 연구에서 검토할 방향으로 기록하였다.'
)

doc.add_heading('6.3 Week 3 세부 계획', level=2)
bullet(
    'Multi-Scale CNN: 다양한 커널 크기(3, 5, 7, 11)로 시간 스케일별 패턴을 뽑는다.'
)
bullet(
    'Multi-Task Learning: 하나의 CNN에서 특성을 뽑은 뒤, '
    '4개 컴포넌트별로 나뉘는 head를 달아 동시에 학습한다.'
)
bullet(
    'Multi-Branch CNN: 100Hz/10Hz/1Hz 센서를 별도 branch로 넣어 '
    '각 해상도에 맞는 특성을 추출한 뒤 합친다.'
)
bullet(
    '목표: 수작업 특성 추출 없이 원시 시계열만으로 학습이 가능한지 확인하고, '
    'ML baseline과의 성능 차이를 정량적으로 비교한다. '
    '딥러닝의 목적은 정확도 향상이 아니라 특성 설계 없이 end-to-end 학습이 가능한지를 검증하는 데 있다.'
)

doc.add_heading('6.4 Week 4 세부 계획', level=2)
bullet(
    'Week 3에서 구현한 딥러닝 모델들의 실험을 완료하고, '
    '하이퍼파라미터 조정을 진행한다.'
)
bullet(
    'ML baseline(LDA, RF, XGBoost 등)과 DL 모델의 성능을 종합적으로 비교 분석한다.'
)
bullet(
    '4주간의 전체 실험 결과를 정리하여 최종 보고서를 작성한다.'
)
body(
    'MFPCA 기반 시계열 형상 분석이나 Transformer 기반 센서 간 attention은 '
    '4주 일정 내에서 다루기 어렵다고 판단하여 이후 연구 과제로 남긴다.'
)


# ═══════════════════════════════════════════════
# 7. 참고문헌
# ═══════════════════════════════════════════════
doc.add_heading('7. 참고문헌', level=1)
refs = [
    'Helwig, N., Pignanelli, E., & Schuetze, A. (2015). Condition monitoring of a complex '
    'hydraulic system using multivariate statistics. IEEE I2MTC.',
    'Chawla, N. V., et al. (2002). SMOTE: Synthetic Minority Over-sampling Technique. JAIR, 16, 321-357.',
    'UCI Machine Learning Repository. Condition monitoring of hydraulic systems. '
    'https://archive.ics.uci.edu/dataset/447/',
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(); r = p.add_run(f'[{i}] {ref}'); sf(r, Pt(10))


# ═══════════════════════════════════════════════
# 저장
# ═══════════════════════════════════════════════
doc.save(OUT_PATH)
print(f"보고서 저장: {OUT_PATH}")
shutil.copy2(OUT_PATH, OUT_COPY)
print(f"복사 완료: {OUT_COPY}")
print(f"파일 크기: {os.path.getsize(OUT_PATH) / 1024:.1f} KB")
