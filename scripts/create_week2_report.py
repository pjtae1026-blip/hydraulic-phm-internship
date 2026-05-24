"""
Week 2 보고서 Word(.docx) 자동 생성 스크립트
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ─── 경로 설정 ───
BASE = os.path.join(os.path.dirname(__file__), '..')
FIG_DIR = os.path.join(BASE, 'reports', 'figures')
OUT_PATH = os.path.join(BASE, 'reports', '2주차_인턴십_보고서.docx')

doc = Document()

# ─── 기본 스타일 설정 ───
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '맑은 고딕'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    hs.font.color.rgb = RGBColor(0, 0, 0)

# 여백 설정
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ─── 헬퍼 함수 ───
def add_title_block():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2주차 인턴십 진행 보고서')
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = '맑은 고딕'

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('유압 시스템 다중결함 동시진단 (UCI Hydraulic Systems)')
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(80, 80, 80)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info = p3.add_run('2026-05-22  |  박정태 (202221215)  |  부산대 기계공학부  |  노유정 교수님 랩')
    info.font.size = Pt(9)
    info.font.color.rgb = RGBColor(100, 100, 100)

    # 구분선
    doc.add_paragraph('─' * 70)


def add_table(headers, rows, col_widths=None):
    """표 생성 헬퍼"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # 데이터
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    return table


def add_figure(filename, caption, width=Inches(5.5)):
    """그림 삽입 + 캡션"""
    path = os.path.join(FIG_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(100, 100, 100)
        r.italic = True
    else:
        doc.add_paragraph(f'[그림 누락: {filename}]')


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10)
        p.add_run(text).font.size = Pt(10)
    else:
        p.add_run(text).font.size = Pt(10)


# ═══════════════════════════════════════════════════════════
# 보고서 본문 작성
# ═══════════════════════════════════════════════════════════

add_title_block()

# ─── 0. 개요 ───
doc.add_heading('0. 2주차 목표 및 개요', level=1)
doc.add_paragraph(
    '2주차의 핵심 목표는 1주차에서 구축한 EDA 기반 위에 (1) 논문(Helwig et al., 2015)의 LDA baseline을 '
    '정밀 재현하고, (2) 다양한 ML 모델과 불균형 처리 기법을 비교하며, (3) 1주차에서 발견한 PS4 센서의 '
    '독자적 거동에 대해 물리적 원인을 규명하는 것이었습니다.'
)
doc.add_paragraph(
    '또한 1주차 보고서에 대한 교수님 피드백을 반영하여, PS4의 물리적 원인을 데이터셋 문서 기반으로 '
    '명확히 하고, 멀티 샘플링 레이트 입력 전략과 향후 계획을 포함하였습니다.'
)

doc.add_paragraph()
doc.add_paragraph('2주차 수행 항목:', style='List Bullet')
add_bullet('Helwig et al. (2015) LDA baseline 정밀 재현 및 논문 목표 달성 확인')
add_bullet('불균형 데이터 처리 4종 비교 실험 (Baseline / class_weight / SMOTE / SMOTE+weight)')
add_bullet('PS4 센서 독자적 거동의 물리적 원인 검증 (MI, 상관분석, PSD)')
add_bullet('ML baseline 5종 비교 (LDA, RF, XGBoost, LightGBM, ExtraTrees)')
add_bullet('멀티 샘플링 레이트 입력 전략 수립')


# ─── 1. LDA Baseline 정밀 재현 ───
doc.add_heading('1. LDA Baseline 정밀 재현', level=1)

doc.add_heading('1.1 실험 설정', level=2)
doc.add_paragraph(
    'Helwig et al. (2015) 논문의 방법론을 재현하였습니다. 17개 센서에서 사이클별 시간도메인 통계량 '
    '11종(mean, std, rms, peak, crest factor, skewness, kurtosis, peak-to-peak, shape factor, '
    'impulse factor, clearance factor)을 추출하여 187차원 tabular feature를 구성하였습니다. '
    'StandardScaler 정규화 후 Linear Discriminant Analysis(LDA)를 적용하고, '
    '5-fold Stratified Cross-Validation으로 평가하였습니다.'
)

doc.add_heading('1.2 결과', level=2)
add_table(
    ['Component', 'Classes', 'Accuracy (%)', 'Macro-F1 (%)', 'MCC', '논문 목표 (%)', '달성'],
    [
        ['Cooler',      '3', '99.95 (±0.09)', '99.93 (±0.12)', '0.9992', '99.0', 'O'],
        ['Valve',       '4', '99.86 (±0.18)', '99.83 (±0.22)', '0.9980', '98.0', 'O'],
        ['Pump',        '3', '99.77 (±0.14)', '99.70 (±0.20)', '0.9963', '95.0', 'O'],
        ['Accumulator', '4', '98.19 (±0.20)', '97.93 (±0.25)', '0.9741', '98.0', 'O'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    '4개 컴포넌트 모두 논문의 목표 정확도를 달성하였습니다. 특히 Cooler(99.95%)와 Valve(99.86%)는 '
    '거의 완벽한 분류 성능을 보이며, 이는 TS1(온도 센서)과 PS1(압력 센서) 안정구간의 높은 분리력에 기인합니다. '
    'Accumulator(98.19%)가 상대적으로 가장 낮은데, 이는 4클래스 분류이며 어큐뮬레이터 열화가 '
    '다른 컴포넌트 신호에 미치는 영향이 미묘하기 때문입니다.'
)

doc.add_heading('1.3 Confusion Matrix', level=2)
add_figure('lda_cm_cooler.png', '[그림 1-1] LDA Confusion Matrix — Cooler (Accuracy: 99.95%)', Inches(3.5))
add_figure('lda_cm_valve.png', '[그림 1-2] LDA Confusion Matrix — Valve (Accuracy: 99.86%)', Inches(3.5))
add_figure('lda_cm_pump.png', '[그림 1-3] LDA Confusion Matrix — Pump (Accuracy: 99.77%)', Inches(3.5))
add_figure('lda_cm_accumulator.png', '[그림 1-4] LDA Confusion Matrix — Accumulator (Accuracy: 98.19%)', Inches(3.5))


# ─── 2. 불균형 데이터 처리 실험 ───
doc.add_heading('2. 불균형 데이터 처리 실험', level=1)

doc.add_heading('2.1 실험 배경', level=2)
doc.add_paragraph(
    'UCI Hydraulic 데이터셋의 클래스 분포를 분석한 결과, Valve(optimal 클래스가 3.1배), '
    'Pump(no_leakage가 2.5배), Accumulator(4클래스 불균형) 등에서 불균형이 존재합니다. '
    '이러한 불균형이 분류 성능, 특히 소수 클래스의 recall에 미치는 영향을 정량화하기 위해 '
    '4종 실험을 설계하였습니다.'
)

doc.add_heading('2.2 실험 설계', level=2)
add_table(
    ['실험', '방법', '설명'],
    [
        ['E1', 'Baseline', 'Random Forest (n=100), 아무 처리 없음'],
        ['E2', 'class_weight', 'RF + class_weight="balanced" (클래스 빈도 역수 가중치)'],
        ['E3', 'SMOTE', 'RF + SMOTE 오버샘플링 (소수 클래스 합성 샘플 생성)'],
        ['E4', 'SMOTE + weight', 'RF + SMOTE + class_weight="balanced" (이중 처리)'],
    ]
)
doc.add_paragraph()
doc.add_paragraph('평가 기준: 5-fold Stratified CV, Accuracy, Macro-F1, MCC, 클래스별 Recall')

doc.add_heading('2.3 주요 결과', level=2)

# Pump 결과 (가장 유의미한 변화)
doc.add_paragraph('Pump (3클래스: no leakage / weak / severe)', style='Heading 3')
add_table(
    ['실험', 'Accuracy (%)', 'Macro-F1 (%)', 'MCC', 'Min Recall (%)'],
    [
        ['E1: Baseline',      '99.73', '99.54', '0.9955', '98.98'],
        ['E2: class_weight',  '99.73', '99.56', '0.9956', '99.39'],
        ['E3: SMOTE',         '99.77', '99.63', '0.9963', '99.39'],
        ['E4: SMOTE+weight',  '99.73', '99.54', '0.9955', '98.98'],
    ]
)

doc.add_paragraph()
doc.add_paragraph('Accumulator (4클래스: optimal / slightly reduced / severely reduced / close to failure)', style='Heading 3')
add_table(
    ['실험', 'Accuracy (%)', 'Macro-F1 (%)', 'MCC', 'Min Recall (%)'],
    [
        ['E1: Baseline',      '98.46', '98.12', '0.9782', '96.24'],
        ['E2: class_weight',  '98.55', '98.25', '0.9798', '96.99'],
        ['E3: SMOTE',         '98.46', '98.14', '0.9786', '96.24'],
        ['E4: SMOTE+weight',  '98.55', '98.25', '0.9796', '96.62'],
    ]
)

doc.add_paragraph()
add_figure('imbalance_macro_f1.png', '[그림 2-1] 불균형 처리 4종 실험 Macro-F1 비교', Inches(5.5))

doc.add_heading('2.4 분석', level=2)
doc.add_paragraph(
    '전반적으로 불균형 처리의 효과가 크지 않았습니다. 이는 187차원 시간도메인 특성이 이미 충분한 '
    '분리력을 갖추고 있어, 클래스 불균형이 있더라도 분류기가 소수 클래스를 잘 학습할 수 있기 때문입니다.'
)
add_bullet('Cooler: ', bold_prefix='')
doc.add_paragraph(
    '  - Cooler: 3클래스가 비교적 균형적이라 4종 실험 간 차이가 미미함',
)
doc.add_paragraph(
    '  - Pump: SMOTE(E3)가 소수 클래스 Min Recall을 98.98% → 99.39%로 개선 (+0.41%p)',
)
doc.add_paragraph(
    '  - Accumulator: class_weight(E2)가 Min Recall을 96.24% → 96.99%로 개선 (+0.75%p)',
)
doc.add_paragraph(
    '  - 결론: 특성 분리가 좋은 데이터에서는 불균형 처리의 marginal gain이 제한적이나, '
    'Accumulator처럼 상대적으로 어려운 타겟에서는 class_weight가 소폭 개선을 제공함',
)


# ─── 3. ML Baseline 5종 비교 ───
doc.add_heading('3. ML Baseline 5종 비교', level=1)

doc.add_heading('3.1 실험 설정', level=2)
doc.add_paragraph(
    'LDA에 더해 4종의 앙상블 모델(Random Forest, XGBoost, LightGBM, Extra Trees)을 비교하였습니다. '
    '모든 모델은 StandardScaler 정규화 후 적용하며, n_estimators=200, 5-fold Stratified CV로 평가합니다.'
)

doc.add_heading('3.2 결과', level=2)
add_table(
    ['Component', 'LDA', 'RF', 'XGBoost', 'LightGBM', 'ExtraTrees'],
    [
        ['Cooler (F1%)',      '99.93', '99.89', '99.84', '99.91', '99.84'],
        ['Valve (F1%)',       '99.83', '99.74', '99.68', '99.74', '99.74'],
        ['Pump (F1%)',        '99.70', '99.59', '99.50', '99.77', '99.59'],
        ['Accumulator (F1%)', '97.93', '98.47', '98.30', '98.64', '98.87'],
    ]
)

doc.add_paragraph()
add_figure('ml_baseline_comparison.png', '[그림 3-1] ML Baseline 5종 모델 Macro-F1 비교 (5-Fold CV)', Inches(5.5))

doc.add_heading('3.3 분석', level=2)
doc.add_paragraph(
    '5개 모델 모두 매우 높은 성능(Macro-F1 > 97.9%)을 보입니다. '
    '주목할 점은 다음과 같습니다:'
)
add_bullet('Cooler, Valve: LDA가 최고 성능. 선형 분리가 최적인 경우 비선형 모델이 추가 이점을 제공하지 못함')
add_bullet('Pump: LightGBM(99.77%)이 LDA(99.70%)를 소폭 상회')
add_bullet('Accumulator: ExtraTrees(98.87%)가 LDA(97.93%)보다 0.94%p 높음. '
           '4클래스 분류의 비선형 경계에서 앙상블 모델이 유리')
add_bullet('전체적으로 187차원 시간도메인 특성의 분리력이 매우 높아, 모델 간 차이가 미미함. '
           '이는 3주차 딥러닝 실험의 동기를 "성능 향상"보다 "수작업 특성 추출 불필요 + '
           '원시 시계열에서 새로운 패턴 발견"으로 설정해야 함을 시사')


# ─── 4. PS4 센서 독자적 거동의 물리적 원인 검증 ───
doc.add_heading('4. PS4 센서 독자적 거동의 물리적 원인 검증', level=1)
doc.add_paragraph(
    '※ 1주차 보고서 교수님 피드백: "PS4의 독자적 거동에 대해 공유된 데이터셋 관련 정보를 찾아 '
    '물리적 원인을 명확히 할 것"'
)

doc.add_heading('4.1 배경', level=2)
doc.add_paragraph(
    '1주차 EDA에서 PS4(압력 센서 4번)가 다른 압력 센서(PS1-PS3, PS5-PS6)와 다른 독자적 거동을 '
    '보이는 것을 발견하였습니다. 이에 대해 3가지 가설을 수립하고 검증하였습니다.'
)
add_bullet('가설 1: PS4가 어큐뮬레이터 분기에 위치하여 accumulator 라벨과 높은 상관을 보임')
add_bullet('가설 2: PS4가 2차 냉각/여과 회로에 위치하여 PS5/PS6과 높은 상관을 보임')
add_bullet('가설 3: PS4가 밸브 하류에 위치하여 valve 라벨과 높은 상관을 보임')

doc.add_heading('4.2 데이터셋 문서 기반 정보', level=2)
doc.add_paragraph(
    'UCI Machine Learning Repository의 Hydraulic Systems 데이터셋 설명(Helwig et al., 2015)에 따르면, '
    '이 유압 시스템은 1차 작동 회로(primary working circuit)와 2차 냉각/여과 회로'
    '(secondary cooling-filtration circuit)로 구성됩니다.'
)
doc.add_paragraph(
    '센서 배치 정보:'
)
add_bullet('PS1 (100Hz): 1차 회로 압력 — 메인 펌프 출력단')
add_bullet('PS2 (100Hz): 1차 회로 압력 — 작동 실린더 측')
add_bullet('PS3 (100Hz): 1차 회로 압력 — 어큐뮬레이터 분기 추정')
add_bullet('PS4 (100Hz): 2차 회로 압력 — 냉각/여과 라인')
add_bullet('PS5 (100Hz): 2차 회로 압력 — 냉각기 입출구')
add_bullet('PS6 (100Hz): 2차 회로 압력 — 리턴 라인')

doc.add_heading('4.3 검증 1: Mutual Information 분석', level=2)
doc.add_paragraph(
    '각 압력 센서의 시간도메인 특성과 4개 결함 라벨 간의 Mutual Information(MI)을 계산하였습니다.'
)
add_figure('ps4_mutual_information.png', '[그림 4-1] 압력 센서별 결함 라벨 Mutual Information 히트맵', Inches(4.5))

doc.add_paragraph('주요 결과:')
add_bullet('PS4-accumulator MI = 0.1255 (PS1-accumulator MI = 0.3435의 약 1/3 수준)')
add_bullet('PS4는 accumulator보다 cooler 라벨과 더 높은 MI를 보임')
add_bullet('→ 가설 1(어큐뮬레이터 분기) 기각: PS4는 어큐뮬레이터와 특별히 강한 연관이 없음')

doc.add_heading('4.4 검증 2: 피어슨 상관계수 분석', level=2)
doc.add_paragraph(
    '6개 압력 센서의 사이클별 평균값 간 피어슨 상관계수를 계산하였습니다.'
)
add_figure('ps4_correlation_matrix.png', '[그림 4-2] 압력 센서 간 상관계수 행렬', Inches(4.5))

doc.add_paragraph('주요 결과:')
add_bullet('PS4-PS1 상관: r = 0.0434 (거의 독립)')
add_bullet('PS4-PS2 상관: r = 0.1284')
add_bullet('PS4-PS3 상관: r = 0.0573')
add_bullet('PS4-PS5 상관: r = 0.7445 (강한 양의 상관)')
add_bullet('PS4-PS6 상관: r = 0.3981')
add_bullet('→ PS4는 1차 회로(PS1-PS3)와 거의 무관하고, 2차 회로(PS5)와 강한 상관')
add_bullet('→ 가설 2(2차 냉각/여과 회로) 지지')

doc.add_heading('4.5 검증 3: PSD(Power Spectral Density) 분석', level=2)
add_figure('ps4_psd_comparison.png', '[그림 4-3] 압력 센서 PSD 비교 (PS1-PS3 vs PS4)', Inches(5))

doc.add_paragraph(
    'PSD 분석 결과, PS1-PS3은 유사한 주파수 특성을 보이는 반면, PS4는 뚜렷하게 다른 스펙트럼 형태를 '
    '나타냅니다. 이는 PS4가 물리적으로 다른 회로에 위치함을 추가적으로 확인합니다.'
)

add_figure('ps4_psd_by_accumulator.png', '[그림 4-4] PS4 PSD — Accumulator 상태별 비교', Inches(4.5))

doc.add_paragraph(
    'Accumulator 상태별 PS4 PSD를 비교하면, accumulator 열화에 따른 PS4 스펙트럼 변화가 관찰되기는 하나, '
    '이는 시스템 전체 압력 맥동의 간접적 영향으로 해석됩니다.'
)

doc.add_heading('4.6 결론', level=2)
doc.add_paragraph(
    '3가지 분석(MI, 상관계수, PSD)을 종합하면, PS4는 2차 냉각/여과 회로에 위치한 센서로 판단됩니다. '
    '1차 작동 회로(PS1-PS3)와 거의 독립적이며, 2차 회로의 PS5와 r=0.7445의 강한 상관을 보입니다. '
    '1주차 보고서에서 관찰한 PS4의 "독자적 거동"은 PS4가 물리적으로 다른 유압 회로에 연결되어 있기 때문이며, '
    '이는 데이터셋 문서의 시스템 구성 정보와 일치합니다.'
)
doc.add_paragraph(
    '이 발견은 향후 딥러닝 모델 설계 시 센서를 회로별로 그룹화하는 multi-branch 구조의 근거가 됩니다.'
)


# ─── 5. 멀티 샘플링 레이트 입력 전략 ───
doc.add_heading('5. 멀티 샘플링 레이트 입력 전략', level=1)
doc.add_paragraph(
    '※ 1주차 보고서 교수님 피드백: "멀티 샘플링 레이트를 학습 모델에 어떻게 입력할지 방안을 고민할 것"'
)

doc.add_heading('5.1 현황', level=2)
doc.add_paragraph(
    'UCI Hydraulic 데이터셋의 17개 센서는 3가지 샘플링 레이트로 수집됩니다:'
)
add_table(
    ['샘플링 레이트', '센서', '사이클당 샘플 수'],
    [
        ['100 Hz', 'PS1-PS6 (압력), EPS1 (모터 전력)', '6,000'],
        ['10 Hz',  'FS1/FS2 (유량), TS1-TS4 (온도)', '600'],
        ['1 Hz',   'VS1 (진동), CE (효율), CP (냉각 전력), SE (안정성 플래그)', '60'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    '2주차까지는 사이클별 통계량(187차원)을 추출하여 샘플링 레이트 차이를 우회하였으나, '
    '3주차 딥러닝에서는 원시 시계열을 직접 입력해야 하므로 전략이 필요합니다.'
)

doc.add_heading('5.2 검토 전략', level=2)

doc.add_paragraph('전략 1: 업샘플링 통일 (Upsample to 100Hz)', style='Heading 3')
add_bullet('모든 센서를 100Hz로 보간(linear interpolation)하여 6,000 timesteps로 통일')
add_bullet('장점: 구현이 간단하고 기존 CNN/LSTM 구조를 그대로 사용 가능')
add_bullet('단점: 1Hz 센서(VS1 등)를 100배 보간하면 인공적 고주파가 생기며, '
           '계산 비용 증가 (17×6,000 = 102,000 입력)')

doc.add_paragraph('전략 2: Multi-Branch CNN (권장)', style='Heading 3')
add_bullet('샘플링 레이트별로 별도 branch를 구성: Branch-100Hz(7ch×6000), Branch-10Hz(6ch×600), Branch-1Hz(4ch×60)')
add_bullet('각 branch에서 독립적으로 특성 추출 후, 공유 fully-connected layer에서 융합')
add_bullet('장점: 각 센서의 원래 해상도를 보존하고, 불필요한 보간 없음')
add_bullet('단점: 구현 복잡도 증가, branch 간 시간 정렬(alignment) 문제')

doc.add_paragraph('전략 3: 통계 특성 + Raw 혼합', style='Heading 3')
add_bullet('고해상도 센서(100Hz)는 raw 시계열 → CNN/LSTM')
add_bullet('저해상도 센서(1Hz, 10Hz)는 통계 특성 → MLP')
add_bullet('두 경로의 출력을 concatenation 후 최종 분류')
add_bullet('장점: 저해상도 센서의 정보 손실 최소화')

doc.add_heading('5.3 3주차 적용 계획', level=2)
doc.add_paragraph(
    '3주차에서는 전략 2(Multi-Branch CNN)를 우선 구현하고, 전략 1(업샘플링 통일)을 '
    'baseline으로 비교할 계획입니다. 전략 2가 가설 H2(multi-rate fusion)의 직접적인 검증이 되며, '
    '4.6절에서 확인한 PS4의 회로별 그룹화 결과를 branch 설계에 반영할 수 있습니다.'
)


# ─── 6. 향후 계획 (Week 3-4) ───
doc.add_heading('6. 향후 계획', level=1)
doc.add_paragraph(
    '※ 1주차 보고서 교수님 피드백: "향후 계획을 다음 자료에 같이 공유할 것"'
)

doc.add_heading('6.1 Week 3 계획: 딥러닝 SOTA 구현', level=2)
add_bullet('MS-CNN (Multi-Scale CNN): 다양한 커널 크기로 다중 시간 스케일 패턴 포착')
add_bullet('Multi-Task Learning: 공유 backbone + 4개 결함별 head (가설 H4 검증)')
add_bullet('Multi-Branch CNN: 샘플링 레이트별 branch 구성 (가설 H2 검증)')
add_bullet('성능 목표: ML baseline 대비 동등 이상의 성능 + 수작업 특성 추출 불필요 입증')

doc.add_heading('6.2 Week 4 계획: Novel Contribution', level=2)
add_bullet('MFPCA (Multivariate Functional PCA) + Transformer 조합 (가설 H3)')
add_bullet('센서 간 attention 메커니즘으로 다중 센서 상호작용 학습 (가설 H5)')
add_bullet('1차/2차 회로 분리 기반의 Hierarchical Multi-Task 구조 제안')
add_bullet('논문 초고 작성 (최종 산출물)')

doc.add_heading('6.3 전체 로드맵', level=2)
add_table(
    ['주차', '핵심 과제', '상태'],
    [
        ['Week 1', 'EDA + LDA Baseline', '완료'],
        ['Week 2', 'ML Baselines + PS4 검증 + 불균형 실험', '완료'],
        ['Week 3', 'DL SOTA (MS-CNN, Multi-Task, Multi-Branch)', '예정'],
        ['Week 4', 'Novel Contribution + 논문 초고', '예정'],
    ]
)


# ─── 7. 참고문헌 ───
doc.add_heading('7. 참고문헌', level=1)
refs = [
    'Helwig, N., Pignanelli, E., & Schuetze, A. (2015). Condition monitoring of a complex hydraulic system using multivariate statistics. IEEE I2MTC.',
    'Chawla, N. V., et al. (2002). SMOTE: Synthetic Minority Over-sampling Technique. JAIR.',
    'UCI Machine Learning Repository. Condition monitoring of hydraulic systems. https://archive.ics.uci.edu/dataset/447/',
]
for i, ref in enumerate(refs, 1):
    doc.add_paragraph(f'[{i}] {ref}')


# ─── 저장 ───
doc.save(OUT_PATH)
print(f"보고서 저장 완료: {OUT_PATH}")
print(f"파일 크기: {os.path.getsize(OUT_PATH) / 1024:.1f} KB")
