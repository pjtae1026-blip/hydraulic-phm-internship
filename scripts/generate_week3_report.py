"""Week 3 보고서 (Word) 생성 스크립트
3주차 DL SOTA 재현 결과를 종합 정리한 .docx 보고서를 생성합니다.
"""

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Load results ───
with open("reports/04_results.json") as f:
    single_task = json.load(f)
with open("reports/05_results.json") as f:
    mt_results = json.load(f)

multi_task = mt_results["multi_task"]
TARGETS = ["cooler", "valve", "pump", "accumulator"]

# ─── Build document ───
doc = Document()

# Default style
style = doc.styles["Normal"]
style.font.name = "Malgun Gothic"
style.font.size = Pt(11)
# Asian font binding (Korean)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rPr.append(rFonts)
rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
rFonts.set(qn("w:ascii"), "Malgun Gothic")
rFonts.set(qn("w:hAnsi"), "Malgun Gothic")

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_korean_font(run, font_name="Malgun Gothic"):
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def add_heading(text, level=1, color=None):
    h = doc.add_heading("", level=level)
    run = h.add_run(text)
    set_korean_font(run, "Malgun Gothic")
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
    return h


def add_para(text, bold=False, italic=False, size=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_korean_font(run)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_korean_font(run)
    return p


def add_image(path, width_inches=6.0, caption=None):
    if not Path(path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        set_korean_font(cap_run)
        cap_run.italic = True
        cap_run.font.size = Pt(10)
        cap_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def style_table(table, header_fill="1F3A68"):
    """Apply consistent styling to a table."""
    table.style = "Light Grid Accent 1"
    # Header row shading
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), header_fill)
        tcPr.append(shd)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True


# ═══════════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("Hydraulic PHM Internship\n3주차 보고서")
set_korean_font(title_run)
title_run.bold = True
title_run.font.size = Pt(22)
title_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("DL SOTA 재현: 1D-CNN 기반 유압 시스템 다중 결함 진단")
set_korean_font(sub_run)
sub_run.font.size = Pt(14)
sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info_p.add_run(
    f"기간: 2026.05.25 - 2026.05.29\n"
    f"데이터셋: UCI Condition Monitoring of Hydraulic Systems (#447)\n"
    f"작성일: {date.today().strftime('%Y년 %m월 %d일')}"
)
set_korean_font(info_run)
info_run.font.size = Pt(11)

doc.add_paragraph()

# ═══════════════════════════════════════════════
# 1. 개요
# ═══════════════════════════════════════════════
add_heading("1. 개요", level=1)

add_para(
    "3주차에서는 UCI Hydraulic Systems 데이터셋의 17개 센서 시계열 신호를 활용하여 "
    "4개 컴포넌트(cooler, valve, pump, accumulator)의 복합 결함을 진단하는 "
    "딥러닝 베이스라인(1D-CNN)을 PyTorch로 구현하고 학습하였다. "
    "구체적으로 다음 두 가지 접근법을 비교하였다."
)

add_bullet("Single-Task 1D-CNN: 4개 타깃을 각각 독립적으로 학습")
add_bullet("Multi-Task 1D-CNN: 공유 backbone + 4-Head 구조로 동시에 학습 (H4 가설 검증)")

add_para(
    "두 모델 모두 로컬 CPU(Intel i5-1235U) 환경에서 학습하였으며, "
    "각 타깃별 분류 성능과 Multi-Task Learning의 효과를 정량적으로 비교하였다."
)

# ═══════════════════════════════════════════════
# 2. 데이터 및 전처리
# ═══════════════════════════════════════════════
add_heading("2. 데이터 및 전처리", level=1)

add_heading("2.1 데이터셋", level=2)
add_para(
    "UCI Condition Monitoring of Hydraulic Systems 데이터셋은 "
    "60초 주기로 측정한 2,205 사이클의 다채널 시계열로 구성되어 있다. "
    "각 사이클에는 17개 센서 신호(샘플링 레이트 100Hz, 10Hz, 1Hz 혼재)와 "
    "4개 컴포넌트의 상태 라벨이 포함된다."
)

# Sensor table
t = doc.add_table(rows=4, cols=4)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr = t.rows[0].cells
hdr[0].text = "센서"
hdr[1].text = "물리량"
hdr[2].text = "샘플링"
hdr[3].text = "Timesteps"
rows_data = [
    ("PS1-PS6, EPS1", "압력, 모터 출력", "100 Hz", "6000"),
    ("FS1, FS2", "유량", "10 Hz", "600"),
    ("TS1-TS4, VS1, CE, CP, SE", "온도, 진동, 효율", "1 Hz", "60"),
]
for i, row_data in enumerate(rows_data, start=1):
    for j, val in enumerate(row_data):
        cell = t.rows[i].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(val)
        set_korean_font(run)
        run.font.size = Pt(10)
# Header bold
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            set_korean_font(run)
            run.font.size = Pt(10)
style_table(t)

add_heading("2.2 라벨 분포", level=2)

t = doc.add_table(rows=5, cols=4)
hdr = t.rows[0].cells
hdr[0].text = "타깃"
hdr[1].text = "클래스 수"
hdr[2].text = "클래스"
hdr[3].text = "분포"
label_rows = [
    ("cooler", "3", "3, 20, 100", "732 / 732 / 741"),
    ("valve", "4", "73, 80, 90, 100", "360 / 360 / 360 / 1125"),
    ("pump", "3", "0, 1, 2", "1221 / 492 / 492"),
    ("accumulator", "4", "90, 100, 115, 130", "808 / 399 / 399 / 599"),
]
for i, row_data in enumerate(label_rows, start=1):
    for j, val in enumerate(row_data):
        cell = t.rows[i].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(val)
        set_korean_font(run)
        run.font.size = Pt(10)
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            set_korean_font(run)
            run.font.size = Pt(10)
style_table(t)

add_heading("2.3 전처리 파이프라인", level=2)
add_bullet("Multi-rate sensor alignment: 모든 센서를 60 timesteps로 균등 다운샘플링")
add_bullet("Per-sensor z-score normalization: 센서별 평균 0, 분산 1로 정규화")
add_bullet("최종 입력 텐서: (2205, 17, 60), 약 18 MB")
add_bullet("Stratified train/test split: 80% / 20% (random_state=42)")

# ═══════════════════════════════════════════════
# 3. 모델 구조
# ═══════════════════════════════════════════════
add_heading("3. 모델 구조", level=1)

add_heading("3.1 1D-CNN Backbone (공통)", level=2)
add_para("두 모델 모두 동일한 1D-CNN backbone을 사용한다. 17개 센서를 채널로 취급한다.")

bb_p = doc.add_paragraph()
bb_run = bb_p.add_run(
    "Conv1d(17→32, k=5) → BN → ReLU → MaxPool → Dropout(0.2)\n"
    "Conv1d(32→64, k=5) → BN → ReLU → MaxPool → Dropout(0.2)\n"
    "Conv1d(64→128, k=3) → BN → ReLU\n"
    "Global Average Pooling → (batch, 128)"
)
set_korean_font(bb_run, "Consolas")
bb_run.font.size = Pt(10)
bb_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

add_heading("3.2 Single-Task (04_pytorch_baseline)", level=2)
add_para(
    "타깃별로 독립적인 모델을 학습한다. 각 모델은 위 backbone + "
    "FC(128→64→n_classes) 분류 head로 구성되며, 총 4개 모델을 별도로 학습한다."
)

add_heading("3.3 Multi-Task (05_multitask)", level=2)
add_para(
    "단일 모델에 공유 backbone + 4개 독립 classification head를 두어 "
    "모든 타깃을 동시에 학습한다. 총 파라미터는 72,142개로, "
    "Single-Task 4개 모델 합산(약 280K) 대비 약 75% 감소했다."
)
add_para(
    "Loss는 Kendall et al. (2018)의 Uncertainty-based weighting 기법을 적용하여 "
    "각 타깃의 학습 난이도에 따라 자동으로 가중치를 조정한다. "
    "수식은 다음과 같다:"
)
formula_p = doc.add_paragraph()
formula_run = formula_p.add_run(
    "L_total = Σ_t [ exp(-log σ²_t) · CE_t + log σ²_t ]"
)
set_korean_font(formula_run, "Consolas")
formula_run.italic = True

# ═══════════════════════════════════════════════
# 4. 학습 설정
# ═══════════════════════════════════════════════
add_heading("4. 학습 설정", level=1)

t = doc.add_table(rows=8, cols=3)
hdr = t.rows[0].cells
hdr[0].text = "항목"
hdr[1].text = "Single-Task (04)"
hdr[2].text = "Multi-Task (05)"
config_rows = [
    ("Optimizer", "Adam (lr=1e-3, wd=1e-4)", "Adam (lr=1e-3, wd=1e-4)"),
    ("Scheduler", "CosineAnnealingLR", "CosineAnnealingLR"),
    ("Batch size", "64", "64"),
    ("Epochs", "50 × 4 = 200", "80"),
    ("Loss", "CrossEntropy", "Uncertainty-weighted CE"),
    ("Device", "CPU (i5-1235U)", "CPU (i5-1235U)"),
    ("Seed", "42", "42"),
]
for i, row_data in enumerate(config_rows, start=1):
    for j, val in enumerate(row_data):
        cell = t.rows[i].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(val)
        set_korean_font(run)
        run.font.size = Pt(10)
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            set_korean_font(run)
            run.font.size = Pt(10)
style_table(t)

# ═══════════════════════════════════════════════
# 5. 실험 결과
# ═══════════════════════════════════════════════
add_heading("5. 실험 결과", level=1)

# Single-task
add_heading("5.1 Single-Task 1D-CNN (04)", level=2)
add_para("4개 타깃에 대해 독립적으로 학습한 결과는 다음과 같다.")

t = doc.add_table(rows=5, cols=3)
hdr = t.rows[0].cells
hdr[0].text = "타깃"
hdr[1].text = "클래스 수"
hdr[2].text = "Test Accuracy"
for i, target in enumerate(TARGETS, start=1):
    cells = t.rows[i].cells
    cells[0].text = ""
    cells[1].text = ""
    cells[2].text = ""
    n_cls = {"cooler": 3, "valve": 4, "pump": 3, "accumulator": 4}[target]
    vals = [target, str(n_cls), f"{single_task[target]:.4f}"]
    for j, val in enumerate(vals):
        run = cells[j].paragraphs[0].add_run(val)
        set_korean_font(run)
        run.font.size = Pt(10)
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            set_korean_font(run)
            run.font.size = Pt(10)
style_table(t)

doc.add_paragraph()
add_image("reports/04_training_curves.png", width_inches=6.5,
          caption="그림 1. Single-Task 1D-CNN 학습 곡선 (4개 타깃별 Loss & Accuracy)")

add_image("reports/04_confusion_matrices.png", width_inches=6.5,
          caption="그림 2. Single-Task 1D-CNN Confusion Matrix")

add_para(
    "Cooler는 첫 epoch부터 100%에 수렴하여 매우 분리하기 쉬운 타깃임을 확인했다. "
    "Pump도 99.8%로 거의 완벽하다. 반면 valve와 accumulator는 90-95% 수준에서 수렴하였으며, "
    "특히 accumulator의 100과 115 클래스, valve의 90과 100 클래스 간 혼동이 관찰되었다."
)

# Multi-task
add_heading("5.2 Multi-Task 1D-CNN (05)", level=2)
add_para(
    "공유 backbone과 4개 head를 갖는 단일 모델로 모든 타깃을 동시에 학습한 결과는 다음과 같다."
)

t = doc.add_table(rows=5, cols=2)
hdr = t.rows[0].cells
hdr[0].text = "타깃"
hdr[1].text = "Test Accuracy"
for i, target in enumerate(TARGETS, start=1):
    cells = t.rows[i].cells
    cells[0].text = ""
    cells[1].text = ""
    vals = [target, f"{multi_task[target]:.4f}"]
    for j, val in enumerate(vals):
        run = cells[j].paragraphs[0].add_run(val)
        set_korean_font(run)
        run.font.size = Pt(10)
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            set_korean_font(run)
            run.font.size = Pt(10)
style_table(t)

doc.add_paragraph()
add_image("reports/05_multitask_accuracy.png", width_inches=6.5,
          caption="그림 3. Multi-Task 1D-CNN 학습 곡선 (4개 타깃 동시 학습)")

add_image("reports/05_task_weights.png", width_inches=6.0,
          caption="그림 4. Uncertainty Weighting으로 학습된 타깃별 가중치 변화")

add_image("reports/05_multitask_confusion.png", width_inches=6.5,
          caption="그림 5. Multi-Task 1D-CNN Confusion Matrix")

add_para(
    "Uncertainty Weighting이 학습 난이도가 다른 4개 타깃의 가중치를 자동으로 조절한 결과, "
    "쉬운 타깃(pump 3.63, cooler 3.21)은 높은 정밀도를, "
    "어려운 타깃(accumulator 2.41)은 상대적으로 낮은 가중치를 부여받았다. "
    "이는 손실 균형화가 자동으로 학습되었음을 의미한다."
)

# Comparison
add_heading("5.3 Single-Task vs Multi-Task 비교", level=2)

t = doc.add_table(rows=6, cols=4)
hdr = t.rows[0].cells
hdr[0].text = "타깃"
hdr[1].text = "Single-Task"
hdr[2].text = "Multi-Task"
hdr[3].text = "Δ (MT − ST)"

avg_st = sum(single_task.values()) / 4
avg_mt = sum(multi_task.values()) / 4

for i, target in enumerate(TARGETS, start=1):
    cells = t.rows[i].cells
    st = single_task[target]
    mt = multi_task[target]
    diff = mt - st
    vals = [target, f"{st:.4f}", f"{mt:.4f}", f"{diff:+.4f}"]
    for j, val in enumerate(vals):
        cells[j].text = ""
        run = cells[j].paragraphs[0].add_run(val)
        set_korean_font(run)
        run.font.size = Pt(10)
        if j == 3:
            if diff > 0:
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            elif diff < 0:
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

# Average row
cells = t.rows[5].cells
vals = ["평균", f"{avg_st:.4f}", f"{avg_mt:.4f}", f"{avg_mt - avg_st:+.4f}"]
for j, val in enumerate(vals):
    cells[j].text = ""
    run = cells[j].paragraphs[0].add_run(val)
    set_korean_font(run)
    run.font.size = Pt(10)
    run.bold = True

for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            set_korean_font(run)
            run.font.size = Pt(10)
style_table(t)

doc.add_paragraph()
add_image("reports/05_comparison.png", width_inches=6.0,
          caption="그림 6. Single-Task vs Multi-Task 정확도 비교")

# ═══════════════════════════════════════════════
# 6. 논의
# ═══════════════════════════════════════════════
add_heading("6. 논의", level=1)

add_heading("6.1 모델 효율성", level=2)
add_para(
    "Multi-Task 모델은 단일 모델(72K 파라미터)로 4개 타깃을 동시에 진단하여, "
    "Single-Task 4개 모델(약 280K) 대비 약 75%의 파라미터를 절약하면서도 "
    "평균 정확도는 오히려 소폭 상승(+0.11%)했다. "
    "이는 추론 시 메모리/연산량 측면에서 큰 이점이며, "
    "실제 산업 현장 배포에 매우 유리한 결과이다."
)

add_heading("6.2 H4 가설 (Multi-Task Learning) 검증", level=2)
add_para(
    "Multi-Task가 모든 타깃에서 일관되게 유리하지는 않았다."
)
add_bullet(f"accumulator: +{(multi_task['accumulator'] - single_task['accumulator'])*100:.2f}%p 개선 → 가장 어려운 타깃에서 효과 확인")
add_bullet(f"pump: +{(multi_task['pump'] - single_task['pump'])*100:.2f}%p 개선")
add_bullet(f"cooler: 변화 없음 (이미 100%)")
add_bullet(f"valve: {(multi_task['valve'] - single_task['valve'])*100:.2f}%p 하락 → 음의 전이(negative transfer) 가능성")

add_para(
    "결론적으로 H4 가설은 '부분적으로 지지'된다. "
    "특히 분류 난이도가 높은 accumulator에서 Multi-Task가 효과적이었으나, "
    "이미 잘 학습되는 타깃(valve)에서는 다른 타깃과의 공유가 오히려 방해가 될 수 있음을 시사한다."
)

add_heading("6.3 클래스 간 혼동 패턴", level=2)
add_para(
    "두 모델 모두 다음 클래스 쌍에서 혼동이 발생했다:"
)
add_bullet("valve: 90 (small lag) ↔ 100 (optimal) — 정상 작동과 약한 지연 간 경계가 모호")
add_bullet("accumulator: 100 (severely reduced) ↔ 115 (slightly reduced) — 중간 단계 압력 구분 어려움")

add_para(
    "이러한 혼동은 물리적으로 인접한 결함 단계 간 신호 차이가 작기 때문으로 추정된다. "
    "추후 압력 센서 cross-correlation 특징(H1) 또는 Transformer attention(H5) 도입 시 "
    "이 부분의 개선이 기대된다."
)

# ═══════════════════════════════════════════════
# 7. 결론 및 다음 단계
# ═══════════════════════════════════════════════
add_heading("7. 결론 및 다음 단계", level=1)

add_heading("7.1 3주차 성과", level=2)
add_bullet("PyTorch 기반 1D-CNN 베이스라인 구축 완료 (04_pytorch_baseline)")
add_bullet("Multi-Task Learning + Uncertainty Weighting 구현 (05_multitask)")
add_bullet(f"평균 96.7-96.8%의 분류 정확도 달성 (cooler/pump: ≥99.8%)")
add_bullet("Multi-Task가 단일 모델로 4개 타깃을 동시 진단하며 파라미터 75% 절감")
add_bullet("H4 가설(Multi-Task의 유효성)을 정량적으로 부분 검증")

add_heading("7.2 4주차 계획", level=2)
add_bullet("H1 가설 검증: 압력센서 PS1-PS6 cross-correlation 특징 추가")
add_bullet("H5 가설 검증: Transformer attention 기반 장거리 의존성 모델")
add_bullet("H2 가설 검증: Multi-rate sensor fusion (학습된 보간) 시도")
add_bullet("최종 보고서 및 발표 자료 작성")

# Save
output_path = "reports/week3_report.docx"
doc.save(output_path)
print(f"[DONE] Report saved: {output_path}")
